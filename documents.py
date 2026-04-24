import os
from datetime import datetime, timezone, timedelta
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils.cell import range_boundaries
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import EXCEL_FILE, WORD_TEMPLATE, ACTS_FOLDER, STATUS_COLORS, FORM_FIELDS
from database import update_excel_row, update_word_path, get_record_by_id


def ensure_dash(value):
    """Возвращает прочерк, если значение пустое"""
    if value is None or str(value).strip() == "":
        return "—"
    return str(value)


def sanitize_filename_part(value, max_len=80):
    """
    Делает строку безопасной для имени файла Windows.
    """
    s = str(value or "").strip()
    if not s:
        return "—"
    # Запрещённые символы Windows: \ / : * ? " < > |
    for ch in '\\/:*?"<>|':
        s = s.replace(ch, "_")
    s = s.replace("\n", " ").replace("\r", " ").replace("\t", " ")
    while "  " in s:
        s = s.replace("  ", " ")
    s = s.strip(" .")
    if len(s) > max_len:
        s = s[:max_len].rstrip(" .")
    return s or "—"


def delete_act_files_for_record(record):
    """
    Удаляет файл(ы) акта для записи.
    1) Если есть точный word_path — удаляем его.
    2) Иначе ищем по номеру акта и шаблонам имён в папке ACTS_FOLDER (в т.ч. старые имена).
    Возвращает количество удалённых файлов.
    """
    removed = 0
    failed = []

    word_path = (record or {}).get("word_path")
    if word_path and os.path.exists(word_path):
        try:
            os.remove(word_path)
            removed += 1
            return removed, failed
        except Exception as e:
            failed.append((word_path, str(e)))

    act_number = str((record or {}).get("act_number") or "").strip()
    if not act_number:
        return removed, failed

    name_part = sanitize_filename_part((record or {}).get("Наименование", ""))
    batch_part = sanitize_filename_part((record or {}).get("№ партии", ""))
    act_part = sanitize_filename_part(act_number)

    # Новый шаблон: Наименование_№партии_№акта.docx
    new_name = f"{name_part}_{batch_part}_{act_part}.docx"

    # Старый шаблон: Акт_{safe_act_number}_{HH-MM-SS}.docx, где П заменяли на P
    safe_act_number = act_number.replace("П", "P")

    for root, _dirs, files in os.walk(ACTS_FOLDER):
        for fn in files:
            if not fn.lower().endswith(".docx"):
                continue

            full = os.path.join(root, fn)

            # Точное совпадение нового имени
            if fn == new_name:
                try:
                    os.remove(full)
                    removed += 1
                except Exception as e:
                    failed.append((full, str(e)))
                continue

            # Старые акты — удаляем по номеру акта в имени
            # Примеры: "Акт_266P_13-10-51.docx" или любые варианты, где встречается номер
            if safe_act_number and safe_act_number in fn:
                try:
                    os.remove(full)
                    removed += 1
                except Exception as e:
                    failed.append((full, str(e)))
                continue

            if act_number in fn:
                try:
                    os.remove(full)
                    removed += 1
                except Exception as e:
                    failed.append((full, str(e)))

    return removed, failed

def get_status_color(status):
    """Получение цвета для статуса"""
    return STATUS_COLORS.get(status, "FFFFFF")

def _norm_header(s):
    """Нормализация заголовка столбца Excel (убирает лишние пробелы)"""
    if s is None:
        return ""
    # Убираем все лишние пробелы (множественные пробелы -> один, табы, переносы строк)
    s = str(s).strip()
    s = ' '.join(s.split())  # Разбиваем по любым whitespace и соединяем одиночным пробелом
    return s

def _get_value(record_data, key):
    v = record_data.get(key, "")
    if v is None:
        return ""
    return str(v).strip()

def _is_meaningful(v):
    s = ("" if v is None else str(v)).strip()
    return bool(s) and s != "—"


def write_to_excel(record_data, record_id, existing_row=None):
    """
    Запись данных в Excel файл с цветовой маркировкой

    Args:
        record_data: словарь с данными записи
        record_id: ID записи в БД
        existing_row: номер существующей строки для обновления (если None - добавляем новую)

    Returns:
        int: номер строки в Excel
    """
    # Открываем книгу
    wb = load_workbook(EXCEL_FILE)
    ws = wb.active

    def find_header_row():
        """
        В некоторых шаблонах заголовки находятся не в первой строке.
        Ищем строку, где в 1-м столбце встречается 'Наименование'.
        """
        for r in range(1, min(ws.max_row, 200) + 1):
            v = ws.cell(row=r, column=1).value
            if isinstance(v, str) and v.strip() == "Наименование":
                return r
        return 1

    header_row = find_header_row()
    data_start_row = header_row + 1

    # Карта "Название столбца" -> индекс колонки (1..n)
    header_map = {}
    for c in range(1, ws.max_column + 1):
        name = _norm_header(ws.cell(row=header_row, column=c).value)
        if name:
            header_map[name] = c

    def set_cell(row, header_name, value):
        # Нормализуем имя заголовка для поиска (чтобы совпадало с ключами header_map)
        norm_name = _norm_header(header_name)
        col = header_map.get(norm_name)
        if not col:
            return
        cell = ws.cell(row=row, column=col, value=value)
        cell.fill = fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        cell.font = Font(name='Arial', size=10)

    def is_row_empty(r, cols_to_check):
        for c in cols_to_check:
            val = ws.cell(row=r, column=c).value
            if val is not None and str(val).strip() != "":
                return False
        return True

    # Определяем номер строки
    if existing_row:
        row_num = existing_row
    else:
        # Находим следующую пустую строку под заголовками.
        # ws.max_row может быть "раздут" форматированием, поэтому идём снизу по данным.
        cols_to_check = list(range(1, 16))  # первые колонки с основными данными
        row_num = None
        for r in range(data_start_row, ws.max_row + 2):
            if is_row_empty(r, cols_to_check):
                row_num = r
                break
        if row_num is None:
            row_num = ws.max_row + 1

    # Применяем стили
    status = record_data.get("status", "")
    fill_color = get_status_color(status)
    fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")

    # Границы ячеек
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Записываем данные по названиям столбцов (чтобы строго соответствовать Excel-шаблону)
    # 1) Номер акта — сохраняем, если такой столбец есть
    act_number = _get_value(record_data, "act_number")
    if act_number:
        set_cell(row_num, "№ акта", act_number)

    # 2) Поля формы — пишем только в те столбцы, которые реально есть в Excel
    for field_name in FORM_FIELDS:
        val = _get_value(record_data, field_name)
        val = ensure_dash(val)  # Пустые значения заменяем на прочерк
        set_cell(row_num, field_name, val)

    # 3) Номер акта в столбец правее таблицы (пустой столбец без заголовка)
    if act_number:
        # Находим последний столбец с заголовком и берем следующий (пустой)
        if header_map:
            last_header_col = max(header_map.values())
            act_col = last_header_col + 1  # Следующий столбец после таблицы
            cell = ws.cell(row=row_num, column=act_col, value=act_number)
            cell.fill = fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            cell.font = Font(name='Arial', size=10)

    # 4) Статус текстом не пишем (цвет строки уже отражает статус)

    # Если на листе есть структурированная Excel-таблица (Table),
    # расширяем её диапазон так, чтобы новая строка попадала "внутрь таблицы".
    # Иначе данные будут записаны, но визуально строка может не появиться в таблице.
    if not existing_row and getattr(ws, "tables", None):
        try:
            # Берём первую таблицу на листе
            table = list(ws.tables.values())[0]
            min_col, min_row, max_col, max_row = range_boundaries(table.ref)
            if row_num > max_row:
                table.ref = f"{ws.cell(row=min_row, column=min_col).coordinate}:{ws.cell(row=row_num, column=max_col).coordinate}"
        except Exception:
            # Даже если расширить таблицу не удалось, сами данные уже записаны в ячейки.
            pass

    # Сохраняем файл
    wb.save(EXCEL_FILE)

    # Обновляем номер строки в БД
    update_excel_row(record_id, row_num)

    return row_num


def delete_excel_row(excel_row):
    """
    Удаляет строку из Excel (со сдвигом вверх).
    """
    if not excel_row:
        return
    wb = load_workbook(EXCEL_FILE)
    ws = wb.active

    ws.delete_rows(int(excel_row), 1)

    # Если на листе есть структурированная таблица, сдвигаем её диапазон,
    # чтобы последняя строка таблицы не "уехала".
    if getattr(ws, "tables", None):
        try:
            table = list(ws.tables.values())[0]
            min_col, min_row, max_col, max_row = range_boundaries(table.ref)
            if int(excel_row) <= max_row:
                table.ref = f"{ws.cell(row=min_row, column=min_col).coordinate}:{ws.cell(row=max_row - 1, column=max_col).coordinate}"
        except Exception:
            pass

    wb.save(EXCEL_FILE)


def create_act_document(record_data, act_number):
    """
    Создание документа акта на основе шаблона Word

    Args:
        record_data: словарь с данными записи
        act_number: номер акта

    Returns:
        str: путь к созданному файлу
    """
    # Проверяем существование шаблона
    if not os.path.exists(WORD_TEMPLATE):
        raise FileNotFoundError(f"Шаблон не найден: {WORD_TEMPLATE}")

    # Открываем шаблон
    doc = Document(WORD_TEMPLATE)

    # Текущая дата и время (МСК, UTC+3)
    msk_tz = timezone(timedelta(hours=3))
    now = datetime.now(msk_tz)
    current_date = now.strftime("%d.%m.%Y")
    current_time = now.strftime("%H:%M:%S")

    # Получаем время проверки из записи или используем текущее
    check_time = record_data.get("check_time", current_time)

    # Функция для правильной замены текста в параграфе
    def set_paragraph_text(paragraph, new_text):
        """Устанавливает текст параграфа, заменяя все runs"""
        # Удаляем все существующие runs
        for run in paragraph.runs:
            r = run._element
            r.getparent().remove(r)
        # Добавляем новый текст
        paragraph.add_run(new_text)

    # Обновляем заголовок (параграф 0)
    if len(doc.paragraphs) > 0:
        set_paragraph_text(doc.paragraphs[0], f"АКТ № {act_number}")

    # Заполняем основные поля (параграфы 3-11 содержат "Наименование:", "Поставщик:" и т.д.)
    field_data = [
        ("Наименование", 3),
        ("Поставщик", 4),
        ("Производитель", 5),
        ("Дата поступления", 6),
        ("Дата проверки", 7),
        ("check_time", 8),  # Специальное поле для времени
        ("№ партии", 9),
        ("Дата изготовления", 10),
        ("Фактическая масса (кг)", 11),
    ]

    for field_name, para_idx in field_data:
        if para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            if field_name == "check_time":
                value = check_time
            else:
                value = ensure_dash(record_data.get(field_name, ""))
            
            # Получаем префикс (текст перед двоеточием)
            old_text = para.text
            if ":" in old_text:
                prefix = old_text.split(":")[0] + ": "
            else:
                prefix = old_text + ": " if old_text else ""
            
            # Устанавливаем новый текст
            set_paragraph_text(para, prefix + str(value))

    # Заполняем таблицу показателей: "Наименование показателя | Норма | Факт | Соответствие"
    # Логика: добавляем строку только если заполнено соответствующее поле "факт"
    
    def find_indicator_table():
        for t in doc.tables:
            try:
                # Теперь принимаем таблицы с 4-5 столбцами (в шаблоне может быть 5)
                if len(t.columns) < 4 or len(t.rows) < 1:
                    continue
                header_cells = [(_norm_header(c.text).lower()) for c in t.rows[0].cells]
                if any("наименование" in x and "показ" in x for x in header_cells) or any("соответ" in x for x in header_cells):
                    return t
            except Exception:
                continue
        return None

    table = find_indicator_table()
    if table is not None and len(table.rows) >= 1:
        # Удаляем все строки кроме шапки (0)
        while len(table.rows) > 1:
            tbl = table._tbl
            tbl.remove(table.rows[1]._tr)

        row_num = 1  # Счетчик для нумерации строк

        # 1) Внешний вид - если заполнен "Внешний вид заявлено"
        vne_vid_zayav = _get_value(record_data, "Внешний вид заявлено")
        vne_vid_fakt = _get_value(record_data, "Внешний вид факт")
        vne_vid_sootv = _get_value(record_data, "Соответствие внешнего вида")
        if _is_meaningful(vne_vid_zayav) or _is_meaningful(vne_vid_fakt):
            row = table.add_row()
            row.cells[0].text = str(row_num)  # № п/п
            row.cells[1].text = "Внешний вид"
            row.cells[2].text = vne_vid_zayav if _is_meaningful(vne_vid_zayav) else ""
            row.cells[3].text = vne_vid_fakt if _is_meaningful(vne_vid_fakt) else ""
            row.cells[4].text = vne_vid_sootv if _is_meaningful(vne_vid_sootv) else ""
            row_num += 1

        # 2) Проверяемые показатели - если заполнен "Проверяемые показатели"
        prover_pokaz = _get_value(record_data, "Проверяемые показатели")
        norm_pasport = _get_value(record_data, "Норматив по паспорту")
        rezult_issled = _get_value(record_data, "Результат исследований")
        zakl_prover = _get_value(record_data, "Заключение по проверяемым показателям")
        if _is_meaningful(prover_pokaz):
            row = table.add_row()
            row.cells[0].text = str(row_num)  # № п/п
            row.cells[1].text = prover_pokaz
            row.cells[2].text = norm_pasport if _is_meaningful(norm_pasport) else ""
            row.cells[3].text = rezult_issled if _is_meaningful(rezult_issled) else ""
            row.cells[4].text = zakl_prover if _is_meaningful(zakl_prover) else ""
            row_num += 1

        # 3) Плотность - если заполнена "Плотность измеренная"
        plotn_izm = _get_value(record_data, "Плотность измеренная г/см³, насыпная плотность кг/м³")
        plotn_pasp = _get_value(record_data, "Плотность по паспорту, кг/м³")
        zakl_plotn = _get_value(record_data, "Заключение по плотности")
        if _is_meaningful(plotn_izm):
            row = table.add_row()
            row.cells[0].text = str(row_num)  # № п/п
            row.cells[1].text = "Плотность измеренная г/см³, насыпная плотность кг/м³"
            row.cells[2].text = plotn_pasp if _is_meaningful(plotn_pasp) else ""
            row.cells[3].text = plotn_izm
            row.cells[4].text = zakl_plotn if _is_meaningful(zakl_plotn) else ""
            row_num += 1

        # 4) Влажность - если заполнена "Влажность измеренная"
        vlazh_izm = _get_value(record_data, "Влажность измеренная, %")
        vlazh_pasp = _get_value(record_data, "Влажность по паспорту, %")
        zakl_vlazh = _get_value(record_data, "Заключение по влажности")
        if _is_meaningful(vlazh_izm):
            row = table.add_row()
            row.cells[0].text = str(row_num)  # № п/п
            row.cells[1].text = "Влажность, %"
            row.cells[2].text = vlazh_pasp if _is_meaningful(vlazh_pasp) else ""
            row.cells[3].text = vlazh_izm
            row.cells[4].text = zakl_vlazh if _is_meaningful(zakl_vlazh) else ""
            row_num += 1

        # 5) Металломагнитные примеси - если заполнены "Метталомагнитные примеси"
        metal_izm = _get_value(record_data, "Метталомагнитные примеси, мг/кг")
        metal_pasp = _get_value(record_data, "Металломагнитные примеси по паспорту, мг/кг")
        zakl_metal = _get_value(record_data, "Заключение по металломагнитным примесям")
        if _is_meaningful(metal_izm):
            row = table.add_row()
            row.cells[0].text = str(row_num)  # № п/п
            row.cells[1].text = "Метталомагнитные примеси, мг/кг"
            row.cells[2].text = metal_pasp if _is_meaningful(metal_pasp) else ""
            row.cells[3].text = metal_izm
            row.cells[4].text = zakl_metal if _is_meaningful(zakl_metal) else ""
            row_num += 1

    # Создаем папку для актов (без разбивки по датам)
    os.makedirs(ACTS_FOLDER, exist_ok=True)

    # Формируем имя файла: Наименование_№партии_№акта.docx
    name_part = sanitize_filename_part(record_data.get("Наименование", ""))
    batch_part = sanitize_filename_part(record_data.get("№ партии", ""))
    act_part = sanitize_filename_part(act_number)
    filename = f"{name_part}_{batch_part}_{act_part}.docx"
    filepath = os.path.join(ACTS_FOLDER, filename)

    # Сохраняем документ
    doc.save(filepath)

    return filepath


def generate_documents(record_data, record_id, status, act_number=None):
    """
    Генерация всех документов для записи

    Args:
        record_data: словарь с данными формы
        record_id: ID записи в БД
        status: статус записи
        act_number: номер акта (если обновление существующей записи)

    Returns:
        dict: информация о созданных документах
    """
    # Получаем полные данные записи из БД
    full_record = get_record_by_id(record_id)
    if not full_record:
        raise ValueError(f"Запись с ID {record_id} не найдена")

    # Добавляем статус в данные
    full_record["status"] = status

    # Определяем номер строки в Excel
    existing_row = full_record.get("excel_row")

    # Если это редактирование существующей записи, удаляем старый акт,
    # чтобы потом создать новый под тем же номером с обновлёнными данными.
    old_word_path = full_record.get("word_path")
    if old_word_path and os.path.exists(old_word_path):
        try:
            os.remove(old_word_path)
        except Exception:
            # Не блокируем сохранение, если старый файл не удалось удалить.
            pass

    # Записываем в Excel
    excel_row = write_to_excel(full_record, record_id, existing_row)

    # Создаем документ Word
    word_path = create_act_document(full_record, full_record["act_number"])
    update_word_path(record_id, word_path)

    return {
        "excel_row": excel_row,
        "word_path": word_path,
        "act_number": full_record["act_number"]
    }
